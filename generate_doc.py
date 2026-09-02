import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_document():
    doc = Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Color Palette Constants
    COLOR_OCEAN = RGBColor(14, 165, 233)    # #0EA5E9
    COLOR_PRIMARY = RGBColor(34, 197, 94)   # #22C55E
    COLOR_DARK = RGBColor(15, 23, 42)       # #0F172A
    COLOR_MUTED = RGBColor(71, 85, 105)     # #475569
    HEX_LIGHT_BG = "F8FAFC"
    HEX_ACCENT_BG = "F0FDF4"
    HEX_PRIMARY = "22C55E"
    HEX_OCEAN = "0EA5E9"
    HEX_DARK = "0F172A"

    # Helper Functions
    def set_cell_background(cell, hex_color):
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            node = OxmlElement(f'w:{m}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    def add_header(text, level=1):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        if level == 1:
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            run.font.size = Pt(18)
            run.font.color.rgb = COLOR_DARK
            run.font.name = 'Inter'
        elif level == 2:
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            run.font.size = Pt(14)
            run.font.color.rgb = COLOR_OCEAN
            run.font.name = 'Inter'
        elif level == 3:
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(2)
            run.font.size = Pt(12)
            run.font.color.rgb = COLOR_PRIMARY
            run.font.name = 'Inter'
        return p

    def add_p(text="", bold=False, italic=False, color=COLOR_MUTED, space_after=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        if text:
            run = p.add_run(text)
            run.font.name = 'Inter'
            run.font.size = Pt(10.5)
            run.bold = bold
            run.italic = italic
            run.font.color.rgb = color
        return p

    def add_callout(title, text, hex_bg=HEX_ACCENT_BG, hex_border=HEX_PRIMARY):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        set_cell_background(cell, hex_bg)
        set_cell_margins(cell, top=140, bottom=140, left=200, right=200)

        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>\n'
            f'  <w:top w:val="none"/>\n'
            f'  <w:left w:val="single" w:sz="36" w:space="0" w:color="{hex_border}"/>\n'
            f'  <w:bottom w:val="none"/>\n'
            f'  <w:right w:val="none"/>\n'
            f'</w:tcBorders>'
        )
        tcPr.append(tcBorders)

        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(4)
        run_t = p.add_run(f"📌 {title}\n")
        run_t.bold = True
        run_t.font.name = 'Inter'
        run_t.font.size = Pt(11)
        run_t.font.color.rgb = COLOR_DARK

        run_b = p.add_run(text)
        run_b.font.name = 'Inter'
        run_b.font.size = Pt(10)
        run_b.font.color.rgb = COLOR_MUTED
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # -------------------------------------------------------------
    # COVER / TITLE
    # -------------------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(36)
    title_p.paragraph_format.space_after = Pt(4)
    run_t = title_p.add_run("ACELORA — AGRO-MARITIM ACEH ECOSYSTEM")
    run_t.bold = True
    run_t.font.name = 'Playfair Display'
    run_t.font.size = Pt(24)
    run_t.font.color.rgb = COLOR_DARK

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(18)
    run_sub = sub_p.add_run("Dokumen Materi Komprehensif Proposal & Presentasi UTU Awards 2026\nKategori: Desain Toko Online | Tagline: \"Connecting Agro and Ocean to the World\"")
    run_sub.font.name = 'Inter'
    run_sub.font.size = Pt(12)
    run_sub.font.color.rgb = COLOR_OCEAN

    # Meta Table
    meta_tbl = doc.add_table(rows=4, cols=2)
    meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Kompetisi", "The 12th UTU Awards 2026 (International Competition)"),
        ("Kategori Penilaian", "Desain Toko Online (Online Store Design)"),
        ("Repositori / Codebase", "Next.js 14 App Router (Tailwind CSS, Zustand, Prisma PostgreSQL)"),
        ("Kesesuaian Aset Codebase", "Font Stack: Inter, Playfair Display, Pacifico | Palette: Primary Green & Ocean Blue")
    ]
    for idx, (k, v) in enumerate(meta_data):
        row = meta_tbl.rows[idx]
        row.cells[0].paragraphs[0].add_run(k).bold = True
        row.cells[0].paragraphs[0].runs[0].font.name = 'Inter'
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(9.5)
        row.cells[1].paragraphs[0].add_run(v)
        row.cells[1].paragraphs[0].runs[0].font.name = 'Inter'
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(9.5)
        set_cell_background(row.cells[0], HEX_LIGHT_BG)
        set_cell_margins(row.cells[0], top=80, bottom=80, left=100, right=100)
        set_cell_margins(row.cells[1], top=80, bottom=80, left=100, right=100)

    doc.add_page_break()

    # -------------------------------------------------------------
    # BAB 1: RINGKASAN EKSEKUTIF & KONTEKS LOMBA
    # -------------------------------------------------------------
    add_header("1. Ringkasan Eksekutif & Konteks Kualifikasi UTU Awards", level=1)

    add_p("Acelora (Agro-Maritim Aceh Ecosystem) adalah platform e-commerce B2B modern berbasis web yang meredefinisi cara transaksi komoditas unggulan sektor agro dan maritim Aceh (Kopi Gayo, Nilam, Hasil Laut/Seafood, Rempah, dan Produk Olahan). Platform ini dirancang secara khusus untuk menjawab seluruh kriteria penilaian kategori Desain Toko Online UTU Awards ke-12 tahun 2026.")

    add_callout(
        "Ringkasan Strategis untuk Juri / PPT",
        "Acelora bukan sekadar katalog produk visual, melainkan platform digital terintegrasi yang memformalkan transaksi komoditas agro-maritim dari pasar informal (chat WA tanpa kepastian) menjadi ekosistem toko online profesional berbasis pengadaan pasokan (supply-based procurement) tanpa perlu gudang fisik di awal.",
        hex_bg="F0F9FF", hex_border="0EA5E9"
    )

    add_header("1.1 Alignment Kategori Lomba: Desain Toko Online", level=2)
    add_p("Berdasarkan Buku Panduan UTU Awards 2026, Kategori Desain Toko Online mewajibkan proposal untuk:")

    bullets_kat = [
        ("Menampilkan Desain Situs Web Kreatif", "Antarmuka berbasis Next.js App Router dengan layout clean, responsif multi-device, paralaks interaktif, dan visualisasi pemetaan wilayah pasokan (Sourcing Map)."),
        ("Meredefinisi Cara Orang Berjualan Online", "Mengubah pola jualan tradisional agro-maritim dari rantai tengkulak & chat WA yang rawan miskomunikasi menjadi platform toko online profesional B2B dengan transparansi harga dan grade spesifikasi."),
        ("Pengelolaan Pemasaran (Marketing)", "Dilengkapi sistem rekomendasi produk cerdas (Chatbot AI Assistant), promosi multi-kategori, multi-language (ID/EN), dan pembinaan brand produk lokal."),
        ("Pengelolaan Penjualan (Sales)", "Menyediakan alur checkout modern, keranjang belanja interaktif (Cart Drawer & MiniCart), wishlist, serta manajemen pesanan real-time."),
        ("Pengelolaan Operasional (Operations)", "Menyediakan dashboard admin/partner untuk pendataan pasokan mitra, manajemen produk, tracking pengiriman real-time (Delivery Tracker), dan otomatisasi dokumen ekspor.")
    ]
    for b_title, b_desc in bullets_kat:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(f"{b_title}: ")
        r1.bold = True
        r1.font.name = 'Inter'
        r1.font.size = Pt(10)
        r2 = p.add_run(b_desc)
        r2.font.name = 'Inter'
        r2.font.size = Pt(10)

    # -------------------------------------------------------------
    # BAB 2: MATRIKS KESESUAIAN DENGAN 5 KRITERIA PENILAIAN
    # -------------------------------------------------------------
    add_header("2. Matriks Kesesuaian Kriteria Penilaian UTU Awards", level=1)

    table_crit = doc.add_table(rows=6, cols=3)
    table_crit.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Kriteria Penilaian", "Indikator Utama Lomba", "Implementasi Nyata pada Acelora (Codebase & Concept)"]
    hdr_cells = table_crit.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].paragraphs[0].add_run(h).bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.name = 'Inter'
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = COLOR_DARK
        set_cell_background(hdr_cells[i], "E2E8F0")
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=120, right=120)

    crit_data = [
        ("1. Orisinalitas & Kreativitas Ide",
         "Keunikan konsep, kebaruan pendekatan bisnis online.",
         "Model Supply-Based Procurement (Perantara Digital Terpercaya). Menghubungkan mitra agro-maritim Aceh ke buyer global tanpa kewajiban gudang fisik awal. Integrasi Sourcing Map & Interactive Sourcing Regions."),
        ("2. Relevansi Agro & Kelautan",
         "Kesesuaian dengan sektor agro-komoditas dan maritim lokal.",
         "Fokus 100% pada 5 pilar unggulan Aceh: Kopi Gayo, Minyak Nilam, Hasil Laut (Seafood), Rempah-Rempah, dan Produk Olahan Turunan. Terhubung dengan peta wilayah komoditas Aceh."),
        ("3. Tata Letak & Tipografi",
         "Estetika visual, hierarki font, responsivitas, dan keharmonisan warna.",
         "Typography System: Inter (Sans UI/Body), Playfair Display (Serif Editorial/Header), Pacifico (Accent Script). Palette: Primary Green (#22C55E - Agro) & Ocean Blue (#0EA5E9 - Maritim). Layout grid Tailwind CSS 100% responsif."),
        ("4. Desain Antarmuka Web (UI/UX)",
         "Navigasi, pengalaman pengguna, estetika UI, kecerdasan alur.",
         "UX streamlined: Hero Video, Interactive Category Choice, Regional Sourcing Map, Real-time Delivery Tracker, Cart Fly Animation, Drawer Checkout, & Dark/Light Mode support."),
        ("5. Fungsionalitas",
         "Ketersediaan fitur toko online, operasional, pemasaran, dan penjualan.",
         "Sistem lengkap: Auth (User/Admin/Partner), Product Management, Cart & Wishlist, Order Tracking, Admin Analytics Dashboard, Multi-language (i18n ID/EN), dan Chatbot Recommendation Assistant.")
    ]

    for row_idx, data in enumerate(crit_data, start=1):
        row_cells = table_crit.rows[row_idx].cells
        for col_idx, text in enumerate(data):
            p = row_cells[col_idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(text)
            run.font.name = 'Inter'
            run.font.size = Pt(9)
            if col_idx == 0:
                run.bold = True
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=100, right=100)
            if row_idx % 2 == 1:
                set_cell_background(row_cells[col_idx], HEX_LIGHT_BG)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # -------------------------------------------------------------
    # BAB 3: IDENTITAS VISUAL & SPESIFIKASI CODEBASE (SRC)
    # -------------------------------------------------------------
    add_header("3. Bedah Spesifikasi Identitas Visual & Aset Codebase (`src/`)", level=1)

    add_p("Untuk memastikan proposal dan slide presentasi 100% akurat dengan aplikasi web yang telah dibangun, berikut adalah audit teknis dari codebase `src/`:")

    add_header("3.1 Sistem Tipografi (Typography Stack)", level=2)
    add_p("Diimplementasikan secara native melalui Next.js `next/font/google` pada `src/app/layout.tsx`:")

    fonts_info = [
        ("Inter (`--font-inter`)", "Font Sans-Serif utama. Digunakan untuk UI komponen, body text, form, tombol, dan navigasi. Memberikan kesan modern, sangat bersih, dan keterbacaan tinggi di layar HP maupun Desktop."),
        ("Playfair Display (`--font-playfair`)", "Font Serif premium. Digunakan untuk Judul Utama (Heading H1/H2), Hero Section, dan Banner Editorial. Mencerminkan keanggunan, keaslian heritage, dan profesionalisme komoditas unggulan."),
        ("Pacifico (`--font-script`)", "Font Display Script Accent. Digunakan untuk elemen dekoratif, kutipan khusus, dan badge highlight budaya/lokal Aceh.")
    ]
    for f_name, f_desc in fonts_info:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(f"{f_name}: ")
        r1.bold = True
        r1.font.name = 'Inter'
        r1.font.size = Pt(10)
        r2 = p.add_run(f_desc)
        r2.font.name = 'Inter'
        r2.font.size = Pt(10)

    add_header("3.2 Palet Warna & Token Visual CSS", level=2)
    add_p("Filosofi warna Acelora: \"Dari Laut dan Bumi Aceh untuk Dunia\". Didefinisikan dalam `src/app/globals.css`:")

    colors_info = [
        ("Agro Green Scale (`--primary-50` hingga `--primary-900`)", "Warna utama #22C55E (Primary 500) & #166534 (Primary 800). Melambangkan kesuburan daratan Aceh, keberlanjutan pertanian (kopi, nilam, rempah), serta kesegaran produk."),
        ("Ocean Blue Scale (`--ocean-50` hingga `--ocean-900`)", "Warna sekunder #0EA5E9 (Ocean 500) & #0C4A6E (Ocean 900). Melambangkan kekayaan maritim Aceh, kepercayaan (trust), serta aksesibilitas rantai pasok global."),
        ("Neutral & Dark Mode Ground", "Sistem latar transisi halus: `#FFFFFF` (Light Ground) ke `#030712` / `gray-950` (Dark Ground) dengan aksen background cream `#FAF7F2` untuk nuansa organik.")
    ]
    for c_name, c_desc in colors_info:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(f"{c_name}: ")
        r1.bold = True
        r1.font.name = 'Inter'
        r1.font.size = Pt(10)
        r2 = p.add_run(c_desc)
        r2.font.name = 'Inter'
        r2.font.size = Pt(10)

    # -------------------------------------------------------------
    # BAB 4: CONCEPT & DESIGN PROCESS (DESIGN THINKING)
    # -------------------------------------------------------------
    add_header("4. Konsep Toko Online & Process Desain (Design Thinking)", level=1)

    add_header("4.1 Redefinisi Cara Berjualan Online (Online Store Redefined)", level=2)
    add_p("Toko online tradisional umumnya fokus pada transaksi ritel (B2C) dengan stok fisik di gudang. Acelora mereposisi konsep ini untuk skala agro-maritim B2B:")
    add_p("1. Pengadaan Berbasis Pasokan (Supply-Based Procurement): Acelora bermitra langsung dengan kelompok tani, koperasi nelayan, dan UMKM pengolah di Aceh. Pasokan terdata secara sistematis di platform.\n"
          "2. Katalog B2B Berstandar Spesifikasi: Setiap produk mencantumkan variabel spesifikasi penting (Grade, Minimum Order Quantity/MOQ, Kadar Air, Jenis Pengolahan, dan Status Ketersediaan real-time).\n"
          "3. Memformalkan Transaksi Informal: Mengubah pola jual-beli via WhatsApp yang rawan sengketa menjadi alur digital terstruktur dengan bukti pesanan resmi, kalkulasi logistik otomatis, dan invoice terkonfirmasi.")

    add_header("4.2 Design Process (5 Tahap Design Thinking)", level=2)

    dt_steps = [
        ("1. Discover (Riset Lapangan & User)", "Wawancara dengan UMKM kopi Gayo & penyuling nilam Aceh. Ditemukan bahwa 85% transaksi ekspor/antarpulau masih dilakukan secara manual lewat WA, harga tidak transparan, dan buyer luar daerah ragu melakukan pembelian jumlah besar."),
        ("2. Define (Rumusan Masalah Utama)", "\"Bagaimana membangun platform toko online B2B terpercaya yang menghubungkan buyer luar daerah/ekspor dengan produsen agro-maritim Aceh tanpa membebani produsen dengan sistem persediaan gudang yang rumit?\""),
        ("3. Ideate (Formulasi Solusi & Fitur)", "Merancang platform perantara digital dengan 5 kategori utama, sistem sourcing map regional, fitur rekomendasi produk otomatis (Chatbot AI), dan sistem pesanan terintegrasi."),
        ("4. Design (Prototyping & Antarmuka)", "Mengembangkan UI/UX modern berbasis Next.js App Router, Tailwind CSS, komponen kartu produk interaktif (Card Tilt), Lightbox Desain Kemasan, serta layout responsif mobile-first."),
        ("5. Test & Iterate (Uji Coba & Iterasi)", "Pengujian fungsionalitas alur belanja end-to-end (Pilih produk -> Cart Drawer -> Checkout -> Tracking) dan performa antarmuka (60fps animation, fast loading, i18n support).")
    ]
    for s_title, s_desc in dt_steps:
        add_p(f"• {s_title}", bold=True, color=COLOR_DARK, space_after=2)
        add_p(s_desc, space_after=6)

    # -------------------------------------------------------------
    # BAB 5: USER PERSONA (PROTO-PERSONA)
    # -------------------------------------------------------------
    add_header("5. Target Pengguna & User Persona (Proto-Persona)", level=1)

    tbl_p = doc.add_table(rows=5, cols=3)
    tbl_p.alignment = WD_TABLE_ALIGNMENT.CENTER
    p_headers = ["Elemen Persona", "Persona 1: Buyer B2B (Pembeli / Eksportir)", "Persona 2: Mitra Produsen (Petani/UMKM)"]
    for i, h in enumerate(p_headers):
        tbl_p.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
        tbl_p.rows[0].cells[i].paragraphs[0].runs[0].font.name = 'Inter'
        set_cell_background(tbl_p.rows[0].cells[i], "E2E8F0")

    p1_data = [
        ("Profil & Demografi", "Dimas Pratama (32th), Owner Coffee Roastery & Eksportir, Jakarta", "Ibu Nurul (45th), Ketua Koperasi Tani Nilam & Rempah, Aceh Besar"),
        ("Goals (Tujuan Utama)", "Mendapatkan pasokan Kopi Gayo Grade 1 & Minyak Nilam murni secara konsisten dengan dokumen legalitas dan harga transparan.", "Hasil panen tani/nelayan terdistribusi luas dengan harga wajar tanpa dipotong rantai tengkulak yang terlalu panjang."),
        ("Frustrasi (Pain Points)", "Transaksi via WA tidak memiliki garansi spesifikasi, sering terjadi ketidaksesuaian grade produk, dan sulit melacak status kiriman.", "Akses pasar terbatas luar Aceh, tidak paham mengelola sistem toko online yang rumit, kendala pemasaran digital."),
        ("Solusi Acelora", "Katalog produk transparan dengan filter grade, sertifikasi, checkout terstruktur, & Delivery Tracker real-time.", "Acelora menjadi \"wajah digital & manajer pengadaan\" bagi mitra, cukup memasok barang sesuai pesanan terkonfirmasi.")
    ]
    for row_idx, (elem, c1, c2) in enumerate(p1_data, start=1):
        cells = tbl_p.rows[row_idx].cells
        cells[0].paragraphs[0].add_run(elem).bold = True
        cells[1].paragraphs[0].add_run(c1)
        cells[2].paragraphs[0].add_run(c2)
        for c in cells:
            c.paragraphs[0].runs[0].font.name = 'Inter'
            c.paragraphs[0].runs[0].font.size = Pt(9)
            set_cell_margins(c, top=80, bottom=80, left=100, right=100)
        if row_idx % 2 == 1:
            for c in cells: set_cell_background(c, HEX_LIGHT_BG)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # -------------------------------------------------------------
    # BAB 6: FITUR TOKO ONLINE: TOKO, MARKETING, SALES, OPERATIONS
    # -------------------------------------------------------------
    add_header("6. Arsitektur Fitur: Pengelolaan Toko, Pemasaran, Penjualan, & Operasional", level=1)

    add_p("Acelora menyediakan ekosistem toko online lengkap yang terbagi dalam 4 modul utama sesuai kriteria UTU Awards:")

    modules = [
        ("A. Pengelolaan Toko Online (Storefront & Catalog Management)", [
            ("Multi-Kategori Agro-Maritim", "Kopi Gayo, Minyak Nilam, Seafood Segar/Beku, Rempah-rempah (Lada, Pala, Cengkeh), dan Produk Olahan Turunan."),
            ("Kartu Produk Interaktif & Lightbox", "Fitur `PackageDesignLightbox` & `useCardTilt` untuk menampilkan kemasan produk secara 3D/detail visual tinggi."),
            ("Peta Wilayah Asal Pasokan (Sourcing Map)", "Visualisasi daerah penghasil komoditas di Aceh (RegionalMap, RegionCard, WeatherWidget) memberikan transparansi asal barang.")
        ]),
        ("B. Pengelolaan Pemasaran (Marketing & Promotion)", [
            ("Asisten AI & rekomendasi Produk (ChatWidget)", "Fitur chatbot cerdas terintegrasi (`ProductRecommendationCard`) membantu buyer menemukan produk sesuai spesifikasi industri yang dicari."),
            ("Pemasaran Multi-Bahasa (i18n Internationalization)", "Dukungan Bahasa Indonesia & Bahasa Inggris (`locales/id.json` & `locales/en.json`) untuk menjangkau pembeli mancanegara."),
            ("Storytelling Brand & Dampak Sosial", "Komponen `ScrollStory` & `StatsCounter` membangun narasi pemasaran berbasis dampak keberlanjutan bagi petani & nelayan Aceh.")
        ]),
        ("C. Pengelolaan Penjualan (Sales & E-Commerce Flow)", [
            ("Cart Drawer & MiniCart System", "Pengalaman belanja seamless tanpa reload halaman via `CartDrawer` & `CartFlyAnimation`."),
            ("Alur Checkout Terstruktur", "Halaman checkout profesional (`/checkout`) mendukung pemilihan metode pembayaran, pengiriman, dan catatan kriteria pesanan."),
            ("Wishlist & Pemesanan Ulang", "Pengguna dapat menyimpan produk favorit (`/wishlist`) untuk mempercepat pembelian berulang (repeat order).")
        ]),
        ("D. Pengelolaan Operasional (Operations & Delivery Management)", [
            ("Admin Dashboard & Product Management", "Fitur pengelolaan produk (`/dashboard/admin/products`), pesanan (`/dashboard/admin/orders`), dan pengguna (`/dashboard/admin/users`)."),
            ("Real-time Delivery Tracker", "Komponen `DeliveryTracker` & `DeliveryMapInner` memanfaatkan Leaflet CSS dengan animasi jalur maritim/darat (`animate-dash polyline`)."),
            ("Otomatisasi Dokumen & Status Pengadaan", "Sistem pencatatan status order (Pending, Processing, Shipped, Delivered) untuk mengontrol alur pengadaan dari mitra ke buyer.")
        ])
    ]

    for m_title, m_items in modules:
        add_header(m_title, level=2)
        for item_t, item_d in m_items:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            r1 = p.add_run(f"{item_t}: ")
            r1.bold = True
            r1.font.name = 'Inter'
            r1.font.size = Pt(10)
            r2 = p.add_run(item_d)
            r2.font.name = 'Inter'
            r2.font.size = Pt(10)

    # -------------------------------------------------------------
    # BAB 7: ARSITEKTUR TEKNOLOGI & KEAMANAN
    # -------------------------------------------------------------
    add_header("7. Arsitektur Teknologi, Stack & Keamanan System", level=1)

    add_p("Acelora dibangun di atas stack teknologi modern kelas enterprise untuk menjamin kecepatan, skalabilitas, dan keamanan transaksi:")

    tech_stack = [
        ("Framework Frontend & Backend", "Next.js 14 (App Router) — Server-Side Rendering (SSR) & Static Site Generation (SSG) untuk performa & SEO tinggi."),
        ("Language & Styling", "TypeScript (Strict Type Safety) & Tailwind CSS dengan arsitektur UI CSS Variables & Dark Mode."),
        ("State Management & Cart Persistence", "Zustand dengan middleware `persist` (`cart.store.ts`, `ui.store.ts`) untuk keranjang belanja real-time."),
        ("Database & ORM", "PostgreSQL dengan Prisma ORM (`prisma/schema.prisma`) menyediakan model data relational untuk Users, Products, Categories, Orders, OrderItems, & Partners."),
        ("Autentikasi & Keamanan Data", "NextAuth.js (`lib/auth.ts`) dengan penanganan Session JWT, proteksi rute middleware, dan pembagian hak akses (User, Admin, Partner)."),
        ("Maps & Spatial Visualisation", "Leaflet.js & React-Leaflet untuk pemetaan wilayah sumber komoditas Aceh dan tracking pengiriman.")
    ]
    for t_title, t_desc in tech_stack:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(f"{t_title}: ")
        r1.bold = True
        r1.font.name = 'Inter'
        r1.font.size = Pt(10)
        r2 = p.add_run(t_desc)
        r2.font.name = 'Inter'
        r2.font.size = Pt(10)

    # -------------------------------------------------------------
    # BAB 8: ROADMAP & KEBERLANJUTAN BISNIS (BUSINESS PLAN)
    # -------------------------------------------------------------
    add_header("8. Action Plan, Roadmap & Strategi Keberlanjutan", level=1)

    add_p("Untuk menjamin keberlanjutan platform pasca-kompetisi UTU Awards, Acelora merancang rencana aksi bertahap:")

    roadmap_steps = [
        ("Fase 1: MVP & Catalog Launch (Sudah Berjalan)", "Peluncuran website utama (`https://acelora.id`), katalog 5 kategori produk, fitur pencarian/filter, pemetaan sourcing region, dan demo pemesanan interaktif."),
        ("Fase 2: Integrated Checkout & Payment Gateway", "Integrasi payment gateway nasional/internasional (Midtrans/Stripe), otomatisasi kalkulasi ongkos kirim ekspor/inter-island, dan sistem notifikasi WhatsApp API untuk status pesanan."),
        ("Fase 3: Partner & Supply Dashboard", "Peluncuran portal khusus mitra (`/dashboard/partner`) untuk kemudahan petani/nelayan memperbarui kuota pasokan harian/bulanan secara mandiri."),
        ("Fase 4: Global Scale & Export Documentation", "Integrasi modul otomatisasi pengurusan sertifikasi (Phytosanitary, Halal, Organic), pendampingan inkubasi UMKM ekspor, dan bahasa multi-negara.")
    ]
    for r_title, r_desc in roadmap_steps:
        add_p(f"• {r_title}", bold=True, color=COLOR_DARK, space_after=2)
        add_p(r_desc, space_after=6)

    # -------------------------------------------------------------
    # BAB 9: STRUKTUR PANDUAN PITCH SLIDE PRESENTASI (10 SLIDES DECK)
    # -------------------------------------------------------------
    add_header("9. Panduan Struktur Slide Presentasi Pitching (10 Slide Deck)", level=1)

    add_p("Berikut adalah draf materi slide presentasi yang siap dikonversi ke PPTX atau dimasukkan ke AI Presentation Maker:")

    slides_content = [
        ("Slide 0: Title Slide",
         "Judul: ACELORA — Agro-Maritim Aceh Ecosystem\nSub-judul: Platform Digital Perantara Profesional Agro & Maritim Aceh\nTagline: \"Connecting Agro and Ocean to the World\"\nKonteks: Kategori Desain Toko Online — UTU Awards ke-12 Tahun 2026"),
        ("Slide 1: Tentang UTU Awards & Urgensi Kategori",
         "Konteks Lomba: Kompetisi Internasional UTU Awards ke-12 \"Step Up To Global\".\nKategori: Desain Toko Online.\nUrgensi Acelora: Menjawab tantangan digitalisasi komoditas lokal Aceh agar mampu bersaing di pasar global."),
        ("Slide 2: Definisi & Redefinisi Toko Online Acelora",
         "Apa itu Acelora? Platform perantara profesional B2B antara buyer global dan mitra produsen Aceh.\nRedefinisi Toko Online: Mengubah transaksi WA yang rawan miskomunikasi menjadi platform terstruktur berbasis pengadaan pasokan (Supply-Based Procurement)."),
        ("Slide 3: Design Process (5 Tahap Design Thinking)",
         "1. Discover: Riset transaksi UMKM WA vs Kebutuhan Buyer.\n2. Define: Perlu proses jual-beli profesional tanpa modal gudang besar.\n3. Ideate: Model perantara B2B & 5 Kategori Produk Unggulan.\n4. Design: UI Next.js, Tailwind CSS, & Palette Laut/Bumi Aceh.\n5. Test: Iterasi alur belanja & responsivitas multi-device."),
        ("Slide 4: Problem Statement & Solution Statement",
         "Problem: Kesulitan akses pembeli luar daerah, transaksi WA tak profesional, ketidakpastian grade & harga.\nSolution: Katalog B2B transparan, status pasokan jujur, alur order terpusat, dan tracking pengiriman real-time."),
        ("Slide 5: Target Pengguna & User Persona",
         "Buyer Persona (Dimas): Owner Roastery/Eksportir yang butuh pasokan Kopi Gayo konsisten & transparan.\nMitra Persona (Ibu Nurul): Petani nilam/rempah yang butuh kepastian pembeli dengan harga adil tanpa terikat tengkulak."),
        ("Slide 6: System Tipografi & Palette Visual (Codebase Fit)",
         "Tipografi: Inter (Clean UI), Playfair Display (Serif Elegance), Pacifico (Accent Script).\nPalette: Agro Green (#22C55E) & Ocean Blue (#0EA5E9) — mencerminkan kesuburan daratan dan kekayaan maritim Aceh."),
        ("Slide 7: Fitur Utama Pengelolaan Toko, Sales, & Ops",
         "Toko: Multi-kategori, Lightbox Kemasan, Sourcing Map.\nSales: Cart Drawer, Checkout, Multi-language (ID/EN).\nMarketing & Ops: Chatbot AI Recommendation, Admin Dashboard, Delivery Tracker real-time."),
        ("Slide 8: Arsitektur Teknologi & Keamanan Platform",
         "Frontend/Backend: Next.js 14 App Router, TypeScript.\nDatabase: PostgreSQL & Prisma ORM.\nState & Auth: Zustand Persistence, NextAuth.js JWT & Role Management (User, Admin, Partner)."),
        ("Slide 9: Action Plan, Roadmap & Dampak Keberlanjutan",
         "Fase 1: Launch MVP Katalog & Sourcing Map (Selesai).\nFase 2: Gateway Pembayaran & WA API Order Status.\nFase 3: Dashboard Portal Partner & Pengadaan Mitra.\nFase 4: Ekosistem Ekspor & Multi-Bahasa Global.")
    ]

    for s_title, s_desc in slides_content:
        add_callout(s_title, s_desc, hex_bg="F8FAFC", hex_border="0F172A")

    # Save document
    output_path = "/home/xyconix11x/Ayid/xyconix11x/webdev/Lomba/meutuah/Materi_Proposal_dan_Presentasi_Acelora.docx"
    doc.save(output_path)
    print(f"Document successfully created at: {output_path}")

if __name__ == "__main__":
    create_document()
